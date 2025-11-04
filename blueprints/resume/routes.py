"""
AI Resume Builder Routes
Intelligent resume creation, optimization, and ATS scanning
"""

from flask import render_template, request, jsonify, send_file, session, flash, redirect, url_for
from flask_login import login_required, current_user
from . import resume_bp
from models import db, Resume, ResumeSection, ResumeTemplate, Job, User
from utils.openai_utils import (
    generate_resume_content,
    optimize_resume_for_job,
    scan_resume_ats,
    generate_cover_letter,
    suggest_resume_improvements
)
from utils.input_validation import sanitize_html, validate_email, validate_url
from extensions import limiter
import json
from datetime import datetime
from io import BytesIO
import pdfkit
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


@resume_bp.route('/')
@login_required
def index():
    """Resume builder dashboard"""
    user_resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.updated_at.desc()).all()
    templates = ResumeTemplate.query.filter_by(is_active=True).all()
    
    return render_template('resume/dashboard.html',
                         resumes=user_resumes,
                         templates=templates,
                         user=current_user)


@resume_bp.route('/create', methods=['GET', 'POST'])
@login_required
@limiter.limit("10 per hour")
def create():
    """Create new resume with AI assistance"""
    if request.method == 'POST':
        data = request.get_json()
        
        # Validate input
        title = sanitize_html(data.get('title', 'My Resume'))
        template_id = data.get('template_id', 1)
        
        # Create resume
        resume = Resume(
            user_id=current_user.id,
            title=title,
            template_id=template_id,
            status='draft'
        )
        
        db.session.add(resume)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'resume_id': resume.id,
            'message': 'Resume created successfully'
        }), 201
    
    # GET request - show creation form
    templates = ResumeTemplate.query.filter_by(is_active=True).all()
    return render_template('resume/create.html', templates=templates)


@resume_bp.route('/<int:resume_id>/edit')
@login_required
def edit(resume_id):
    """Edit existing resume"""
    resume = Resume.query.get_or_404(resume_id)
    
    # Check ownership
    if resume.user_id != current_user.id:
        flash('You do not have permission to edit this resume', 'error')
        return redirect(url_for('resume.index'))
    
    templates = ResumeTemplate.query.filter_by(is_active=True).all()
    jobs = Job.query.filter_by(status='active').order_by(Job.created_at.desc()).limit(50).all()
    
    return render_template('resume/edit.html',
                         resume=resume,
                         templates=templates,
                         jobs=jobs,
                         user=current_user)


@resume_bp.route('/<int:resume_id>/section', methods=['POST'])
@login_required
@limiter.limit("20 per hour")
def add_section(resume_id):
    """Add section to resume"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    
    section = ResumeSection(
        resume_id=resume_id,
        section_type=data.get('type'),
        title=sanitize_html(data.get('title', '')),
        content=sanitize_html(data.get('content', '')),
        order=data.get('order', 0)
    )
    
    db.session.add(section)
    resume.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({
        'success': True,
        'section_id': section.id,
        'message': 'Section added successfully'
    }), 201


@resume_bp.route('/<int:resume_id>/section/<int:section_id>', methods=['PUT', 'DELETE'])
@login_required
@limiter.limit("30 per hour")
def manage_section(resume_id, section_id):
    """Update or delete resume section"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    section = ResumeSection.query.get_or_404(section_id)
    
    if section.resume_id != resume_id:
        return jsonify({'error': 'Section does not belong to this resume'}), 400
    
    if request.method == 'DELETE':
        db.session.delete(section)
        resume.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'message': 'Section deleted'}), 200
    
    # PUT request - update section
    data = request.get_json()
    
    if 'title' in data:
        section.title = sanitize_html(data['title'])
    if 'content' in data:
        section.content = sanitize_html(data['content'])
    if 'order' in data:
        section.order = data['order']
    
    section.updated_at = datetime.utcnow()
    resume.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': 'Section updated successfully'
    }), 200


@resume_bp.route('/<int:resume_id>/ai-generate', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def ai_generate(resume_id):
    """Generate resume content using AI"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    section_type = data.get('section_type')
    context = data.get('context', {})
    
    # Add user context
    context['user'] = {
        'name': current_user.full_name,
        'email': current_user.email,
        'major': getattr(current_user, 'major', ''),
        'graduation_year': getattr(current_user, 'graduation_year', ''),
        'role': current_user.role
    }
    
    try:
        # Generate content using OpenAI
        generated_content = generate_resume_content(section_type, context)
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'message': 'Content generated successfully'
        }), 200
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to generate content',
            'details': str(e)
        }), 500


@resume_bp.route('/<int:resume_id>/optimize', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def optimize_for_job(resume_id):
    """Optimize resume for specific job posting using AI"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    job_id = data.get('job_id')
    
    if not job_id:
        return jsonify({'error': 'Job ID required'}), 400
    
    job = Job.query.get_or_404(job_id)
    
    # Get resume content
    resume_content = {
        'title': resume.title,
        'sections': []
    }
    
    for section in resume.sections:
        resume_content['sections'].append({
            'type': section.section_type,
            'title': section.title,
            'content': section.content
        })
    
    try:
        # Optimize using AI
        optimized = optimize_resume_for_job(resume_content, job)
        
        return jsonify({
            'success': True,
            'optimized_content': optimized['content'],
            'suggestions': optimized['suggestions'],
            'keyword_match': optimized['keyword_match'],
            'message': 'Resume optimized successfully'
        }), 200
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to optimize resume',
            'details': str(e)
        }), 500


@resume_bp.route('/<int:resume_id>/ats-scan', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def ats_scan(resume_id):
    """Scan resume for ATS compatibility"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Get resume content
    resume_content = {
        'title': resume.title,
        'sections': []
    }
    
    for section in resume.sections:
        resume_content['sections'].append({
            'type': section.section_type,
            'title': section.title,
            'content': section.content
        })
    
    try:
        # Scan using AI
        scan_results = scan_resume_ats(resume_content)
        
        return jsonify({
            'success': True,
            'score': scan_results['score'],
            'issues': scan_results['issues'],
            'recommendations': scan_results['recommendations'],
            'keyword_density': scan_results['keyword_density'],
            'formatting_score': scan_results['formatting_score'],
            'message': 'ATS scan completed'
        }), 200
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to scan resume',
            'details': str(e)
        }), 500


@resume_bp.route('/<int:resume_id>/cover-letter', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def generate_cover_letter_route(resume_id):
    """Generate cover letter for job application"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    job_id = data.get('job_id')
    
    if not job_id:
        return jsonify({'error': 'Job ID required'}), 400
    
    job = Job.query.get_or_404(job_id)
    
    # Get resume content
    resume_content = {
        'user': {
            'name': current_user.full_name,
            'email': current_user.email,
            'major': getattr(current_user, 'major', ''),
            'graduation_year': getattr(current_user, 'graduation_year', '')
        },
        'sections': []
    }
    
    for section in resume.sections:
        resume_content['sections'].append({
            'type': section.section_type,
            'title': section.title,
            'content': section.content
        })
    
    try:
        # Generate cover letter using AI
        cover_letter = generate_cover_letter(resume_content, job)
        
        return jsonify({
            'success': True,
            'cover_letter': cover_letter,
            'message': 'Cover letter generated successfully'
        }), 200
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to generate cover letter',
            'details': str(e)
        }), 500


@resume_bp.route('/<int:resume_id>/suggestions', methods=['GET'])
@login_required
@limiter.limit("10 per hour")
def get_suggestions(resume_id):
    """Get AI-powered improvement suggestions"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Get resume content
    resume_content = {
        'title': resume.title,
        'sections': []
    }
    
    for section in resume.sections:
        resume_content['sections'].append({
            'type': section.section_type,
            'title': section.title,
            'content': section.content
        })
    
    try:
        # Get suggestions using AI
        suggestions = suggest_resume_improvements(resume_content)
        
        return jsonify({
            'success': True,
            'suggestions': suggestions,
            'message': 'Suggestions generated'
        }), 200
        
    except Exception as e:
        return jsonify({
            'error': 'Failed to generate suggestions',
            'details': str(e)
        }), 500


@resume_bp.route('/<int:resume_id>/export/<format>')
@login_required
def export(resume_id, format):
    """Export resume in various formats (PDF, DOCX, TXT)"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        flash('You do not have permission to export this resume', 'error')
        return redirect(url_for('resume.index'))
    
    template = ResumeTemplate.query.get(resume.template_id)
    
    if format == 'pdf':
        return export_pdf(resume, template)
    elif format == 'docx':
        return export_docx(resume, template)
    elif format == 'txt':
        return export_txt(resume)
    else:
        flash('Invalid export format', 'error')
        return redirect(url_for('resume.edit', resume_id=resume_id))


def export_pdf(resume, template):
    """Export resume as PDF"""
    try:
        # Render HTML template
        html_content = render_template('resume/export_pdf.html',
                                      resume=resume,
                                      template=template)
        
        # Convert to PDF
        pdf = pdfkit.from_string(html_content, False)
        
        # Create response
        response = BytesIO(pdf)
        filename = f"{resume.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return send_file(
            response,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        flash(f'Failed to export PDF: {str(e)}', 'error')
        return redirect(url_for('resume.edit', resume_id=resume.id))


def export_docx(resume, template):
    """Export resume as Word document"""
    try:
        doc = Document()
        
        # Set up styling
        style = doc.styles['Normal']
        font = style.font
        font.name = template.font_family if template else 'Calibri'
        font.size = Pt(11)
        
        # Header with name
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(current_user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        
        # Contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = f"{current_user.email}"
        if hasattr(current_user, 'phone') and current_user.phone:
            contact_text += f" | {current_user.phone}"
        if hasattr(current_user, 'linkedin_url') and current_user.linkedin_url:
            contact_text += f" | {current_user.linkedin_url}"
        contact.add_run(contact_text)
        
        doc.add_paragraph()  # Spacing
        
        # Add sections
        for section in sorted(resume.sections, key=lambda x: x.order):
            # Section title
            section_title = doc.add_paragraph()
            section_run = section_title.add_run(section.title.upper())
            section_run.bold = True
            section_run.font.size = Pt(12)
            
            # Add horizontal line
            doc.add_paragraph('_' * 80)
            
            # Section content
            content_para = doc.add_paragraph(section.content)
            content_para.paragraph_format.space_after = Pt(12)
        
        # Save to BytesIO
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        filename = f"{resume.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        flash(f'Failed to export DOCX: {str(e)}', 'error')
        return redirect(url_for('resume.edit', resume_id=resume.id))


def export_txt(resume):
    """Export resume as plain text"""
    try:
        content = f"{current_user.full_name}\n"
        content += f"{current_user.email}"
        
        if hasattr(current_user, 'phone') and current_user.phone:
            content += f" | {current_user.phone}"
        
        content += "\n\n"
        
        # Add sections
        for section in sorted(resume.sections, key=lambda x: x.order):
            content += f"{section.title.upper()}\n"
            content += f"{'=' * len(section.title)}\n"
            content += f"{section.content}\n\n"
        
        # Create response
        txt_file = BytesIO(content.encode('utf-8'))
        filename = f"{resume.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"
        
        return send_file(
            txt_file,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        flash(f'Failed to export TXT: {str(e)}', 'error')
        return redirect(url_for('resume.edit', resume_id=resume.id))


@resume_bp.route('/<int:resume_id>/delete', methods=['POST'])
@login_required
def delete(resume_id):
    """Delete resume"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    db.session.delete(resume)
    db.session.commit()
    
    flash('Resume deleted successfully', 'success')
    return jsonify({'success': True}), 200


@resume_bp.route('/templates')
@login_required
def templates():
    """View available resume templates"""
    all_templates = ResumeTemplate.query.filter_by(is_active=True).all()
    
    return render_template('resume/templates.html', templates=all_templates)


@resume_bp.route('/<int:resume_id>/share', methods=['POST'])
@login_required
def share(resume_id):
    """Generate shareable link for resume"""
    resume = Resume.query.get_or_404(resume_id)
    
    if resume.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json()
    is_public = data.get('is_public', False)
    
    resume.is_public = is_public
    
    if is_public:
        # Generate share token if not exists
        if not resume.share_token:
            import secrets
            resume.share_token = secrets.token_urlsafe(32)
    
    db.session.commit()
    
    share_url = url_for('resume.public_view', token=resume.share_token, _external=True) if is_public else None
    
    return jsonify({
        'success': True,
        'is_public': is_public,
        'share_url': share_url
    }), 200


@resume_bp.route('/view/<token>')
def public_view(token):
    """Public view of shared resume"""
    resume = Resume.query.filter_by(share_token=token, is_public=True).first_or_404()
    template = ResumeTemplate.query.get(resume.template_id)
    
    return render_template('resume/public_view.html',
                         resume=resume,
                         template=template,
                         user=resume.user)
